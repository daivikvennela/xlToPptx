# block_replacer.py

import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
import base64

def load_block_template(filename):
    path = os.path.join('templates', 'blocks', filename)
    with open(path, 'r') as f:
        return f.read()

def embedImage(doc: Document, image_data: str, placeholder: str = '[EXHIBIT_A_IMAGE_1]'):
    """
    Embed an image into a DOCX document at the location of a placeholder.
    
    Args:
        doc: The DOCX document object
        image_data: Base64 encoded image string
        placeholder: The placeholder text to replace with the image
    
    Returns:
        bool: True if image was successfully embedded, False otherwise
    """
    try:
        print(f"[DEBUG] Starting image embedding for placeholder: {placeholder}")
        
        # Validate input
        if not image_data or not isinstance(image_data, str):
            print("[ERROR] Invalid image data: must be non-empty string")
            return False
        
        # Decode base64 image data
        try:
            image_bytes = base64.b64decode(image_data)
            print(f"[DEBUG] Decoded base64 image data, size: {len(image_bytes)} bytes")
        except Exception as e:
            print(f"[ERROR] Failed to decode base64 image data: {str(e)}")
            return False
        
        # Validate minimum size
        if len(image_bytes) < 8:
            print("[ERROR] Image data too small to be valid")
            return False
        
        # Validate PNG header
        if not image_bytes.startswith(b'\x89PNG\r\n\x1a\n'):
            print("[ERROR] Invalid PNG header")
            return False
        
        # Open and validate image with Pillow
        try:
            image = Image.open(io.BytesIO(image_bytes))
            print(f"[DEBUG] Opened image: format={image.format}, size={image.size}, mode={image.mode}")
            
            # Convert to RGB if necessary (for PNG with transparency)
            if image.mode in ('RGBA', 'LA', 'P'):
                # Create white background for transparent images
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                image = background
                print("[DEBUG] Converted transparent image to RGB with white background")
            elif image.mode != 'RGB':
                image = image.convert('RGB')
                print(f"[DEBUG] Converted image from {image.mode} to RGB")
        except Exception as e:
            print(f"[ERROR] Failed to process image with Pillow: {str(e)}")
            print(f"[DEBUG] Image bytes size: {len(image_bytes)}")
            print(f"[DEBUG] First 100 bytes: {image_bytes[:100]}")
            return False
        
        # Resize image to reasonable dimensions (max width 6 inches)
        max_width_inches = 6.0
        max_width_pixels = int(max_width_inches * 96)  # 96 DPI for screen
        
        original_size = image.size
        if image.width > max_width_pixels:
            ratio = max_width_pixels / image.width
            new_width = max_width_pixels
            new_height = int(image.height * ratio)
            image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            print(f"[DEBUG] Resized image from {original_size} to {image.size}")
        
        # Convert back to bytes
        img_byte_arr = io.BytesIO()
        image.save(img_byte_arr, format='PNG', optimize=True)
        img_byte_arr = img_byte_arr.getvalue()
        print(f"[DEBUG] Converted image to PNG format, size: {len(img_byte_arr)} bytes")
        
        # Find and replace placeholder in document
        found_placeholder = False
        placeholder_count = 0
        
        def process_paragraph(paragraph):
            nonlocal found_placeholder, placeholder_count
            if placeholder in paragraph.text:
                placeholder_count += 1
                print(f"[DEBUG] Found placeholder '{placeholder}' in paragraph #{placeholder_count}")
                
                # Clear the paragraph and add centered image
                paragraph.clear()
                run = paragraph.add_run()
                
                # Calculate image width in inches
                width_inches = min(image.width / 96, max_width_inches)
                
                # Add image to the run
                run.add_picture(io.BytesIO(img_byte_arr), width=Inches(width_inches))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                found_placeholder = True
                print(f"[DEBUG] Successfully embedded image in paragraph #{placeholder_count}")
        
        def process_table(table):
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)
        
        def process_block(block):
            for paragraph in block.paragraphs:
                process_paragraph(paragraph)
            for table in getattr(block, 'tables', []):
                process_table(table)
        
        # Process main document
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph)
        
        for table in doc.tables:
            process_table(table)
        
        # Process headers and footers
        for section in doc.sections:
            process_block(section.header)
            process_block(section.footer)
        
        # Process footnotes
        if hasattr(doc, 'part') and hasattr(doc.part, 'footnotes'):
            for footnote in doc.part.footnotes.part.footnotes:
                for paragraph in footnote.paragraphs:
                    process_paragraph(paragraph)
        
        if not found_placeholder:
            print(f"[WARNING] Placeholder '{placeholder}' not found in document")
            print(f"[DEBUG] Document has {len(doc.paragraphs)} paragraphs")
            return False
        
        print(f"[DEBUG] Image embedding completed successfully. Found {placeholder_count} placeholder(s)")
        return True
        
    except Exception as e:
        print(f"[ERROR] Critical error in image embedding: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def generate_signature_block(grantor_name, trust_entity_name=None, name=None, title=None, block_type='individual', state=None, county=None, name_of_individuals=None, type_of_authority=None, instrument_for=None):
    if block_type == 'individual':
        template = load_block_template('individual_signature.txt')
        return template.replace('[Grantor Name]', grantor_name or '')
    else:
        template = load_block_template('entity_signature.txt')
        return template.replace('[Trust/Entity Name]', trust_entity_name or '') \
                      .replace('[Name]', name or '') \
                      .replace('[Title]', title or '')

def generate_enhanced_signature_block(owner_type, grantor_name=None, trust_entity_name=None, name=None, title=None, state=None, num_signatures=1):
    """Generate enhanced signature block based on owner type following married_couple_signature model"""
    template_mapping = {
        'individual': 'individual_signature_enhanced.txt',
        'corporation': 'corporation_signature_enhanced.txt', 
        'llc': 'llc_signature_enhanced.txt',
        'lp': 'lp_signature_enhanced.txt',
        'married_couple': 'married_couple_signature_enhanced.txt',
        'sole_owner_married_couple': 'sole_owner_married_couple_enhanced.txt'
    }
    
    template_file = template_mapping.get(owner_type, 'individual_signature_enhanced.txt')
    
    try:
        template_path = os.path.join('templates', 'sigBlocks', template_file)
        with open(template_path, 'r', encoding='utf-8') as f:
            template = f.read()
    except FileNotFoundError:
        # Fallback to basic template
        if owner_type == 'individual':
            template = load_block_template('individual_signature.txt')
        else:
            template = load_block_template('entity_signature.txt')
    
    # Replace placeholders based on owner type
    result = template
    if grantor_name:
        result = result.replace('[Grantor Name]', grantor_name)
    if trust_entity_name:
        result = result.replace('[Trust/Entity Name]', trust_entity_name)
    if name:
        result = result.replace('[Name]', name)
    if title:
        result = result.replace('[Title]', title)
    if state:
        result = result.replace('[State]', state)
    
    # Handle multiple signatures for certain types
    if owner_type == 'married_couple' and num_signatures == 2:
        # The template already has placeholders for husband and wife
        pass
    elif num_signatures > 1 and owner_type not in ['married_couple', 'sole_owner_married_couple']:
        # Add additional signature lines for other types
        signature_lines = []
        for i in range(num_signatures):
            signature_lines.append(f"\nBy:________________________\n[Name {i+1}]\n[Title {i+1}]")
        
        # Replace the single signature block with multiple
        single_sig_pattern = "By:________________________\n[Name]\n[Title]"
        if single_sig_pattern in result:
            result = result.replace(single_sig_pattern, "\n".join(signature_lines))
    
    return result

def generate_notary_block(state, county, name_of_individuals, type_of_authority=None, instrument_for=None, block_type='individual'):
    if block_type == 'individual':
        template = load_block_template('individual_notary.txt')
        return template.replace('[State]', state or '') \
                      .replace('[County]', county or '') \
                      .replace('[NAME(S) OF INDIVIDUAL(S)]', name_of_individuals or '')
    else:
        template = load_block_template('entity_notary.txt')
        return template.replace('[State]', state or '') \
                      .replace('[County]', county or '') \
                      .replace('[NAME(S) OF INDIVIDUAL(S)]', name_of_individuals or '') \
                      .replace('[TYPE OF AUTHORITY]', type_of_authority or '') \
                      .replace('[NAME OF ENTITY OR TRUST WHOM INSTRUMENT WAS EXECUTED FOR]', instrument_for or '')

def generate_enhanced_combined_block(owner_type, grantor_name=None, trust_entity_name=None, name=None, title=None, 
                                    state=None, county=None, name_of_individuals=None, type_of_authority=None, 
                                    instrument_for=None, num_signatures=1, include_signature=True, include_notary=True, 
                                    embed_notary_in_signature=True):
    """Generate combined signature and notary block with embedding logic"""
    result = {}
    
    # Generate signature block
    if include_signature:
        signature_block = generate_enhanced_signature_block(
            owner_type, grantor_name, trust_entity_name, name, title, state, num_signatures
        )
        result['signature_block'] = signature_block
    
    # Generate notary block
    if include_notary:
        block_type = 'individual' if owner_type == 'individual' else 'entity'
        notary_block = generate_notary_block(state, county, name_of_individuals, type_of_authority, instrument_for, block_type)
        result['notary_block'] = notary_block
    
    # Generate combined block
    if include_signature and include_notary and embed_notary_in_signature:
        # Embed notary block within signature block
        combined_block = signature_block.replace('[Notary Block]', notary_block)
        result['combined_block'] = combined_block
    elif include_signature and include_notary and not embed_notary_in_signature:
        # Show blocks separately
        result['combined_block'] = f"{signature_block}\n\n\n{notary_block}"
    elif include_signature:
        # Remove [Notary Block] placeholder
        result['combined_block'] = signature_block.replace('[Notary Block]', '').strip()
    elif include_notary:
        result['combined_block'] = notary_block
    
    # Generate step breakdown
    step_breakdown = []
    if include_notary:
        step_breakdown.append("STEP 1: Notary Block Configuration")
        step_breakdown.append(f"  - State: {state or 'Not specified'}")
        step_breakdown.append(f"  - County: {county or 'Not specified'}")
        step_breakdown.append(f"  - Individual(s): {name_of_individuals or 'Not specified'}")
        if type_of_authority:
            step_breakdown.append(f"  - Authority: {type_of_authority}")
        if instrument_for:
            step_breakdown.append(f"  - Instrument For: {instrument_for}")
        step_breakdown.append("")
    
    if include_signature:
        step_breakdown.append("STEP 2: Signature Block Configuration")
        step_breakdown.append(f"  - Owner Type: {owner_type.replace('_', ' ').title()}")
        step_breakdown.append(f"  - Number of Signatures: {num_signatures}")
        if grantor_name:
            step_breakdown.append(f"  - Grantor Name: {grantor_name}")
        if trust_entity_name:
            step_breakdown.append(f"  - Trust/Entity Name: {trust_entity_name}")
        if name:
            step_breakdown.append(f"  - Signatory Name: {name}")
        if title:
            step_breakdown.append(f"  - Title: {title}")
        step_breakdown.append("")
    
    if embed_notary_in_signature and include_signature and include_notary:
        step_breakdown.append("EMBEDDING: Notary Block embedded within Signature Block")
    elif include_signature and include_notary:
        step_breakdown.append("LAYOUT: Signature and Notary Blocks displayed separately")
    
    result['step_breakdown'] = "\n".join(step_breakdown)
    result['embed_notary_in_signature'] = embed_notary_in_signature
    result['include_signature'] = include_signature
    result['include_notary'] = include_notary
    
    return result

def get_all_block_previews(grantor_name, trust_entity_name, name, title, state, county, name_of_individuals, type_of_authority, instrument_for):
    preview = {
        'individual_signature': generate_signature_block(grantor_name, block_type='individual'),
        'individual_notary': generate_notary_block(state, county, name_of_individuals, block_type='individual'),
        'entity_signature': generate_signature_block(grantor_name, trust_entity_name, name, title, block_type='entity'),
        'entity_notary': generate_notary_block(state, county, name_of_individuals, type_of_authority, instrument_for, block_type='entity'),
    }
    return preview

def replace_signature_and_notary_blocks(doc: Document, mapping: dict):
    # Determine party type
    grantee_type = mapping.get('[Grantee Type]', '').strip().lower()
    is_individual = grantee_type == 'individual'
    # Prepare values for template filling
    grantor_name = mapping.get('[Grantor Name]', '')
    trust_entity_name = mapping.get('[Trust/Entity Name]', '')
    name = mapping.get('[Name]', '')
    title = mapping.get('[Title]', '')
    state = mapping.get('[State]', '')
    county = mapping.get('[County]', '')
    name_of_individuals = mapping.get('[NAME(S) OF INDIVIDUAL(S)]', '')
    type_of_authority = mapping.get('[TYPE OF AUTHORITY]', '')
    instrument_for = mapping.get('[NAME OF ENTITY OR TRUST WHOM INSTRUMENT WAS EXECUTED FOR]', '')
    # Generate blocks
    if is_individual:
        sig_block = generate_signature_block(grantor_name, block_type='individual')
        notary_block = generate_notary_block(state, county, name_of_individuals, block_type='individual')
    else:
        sig_block = generate_signature_block(grantor_name, trust_entity_name, name, title, block_type='entity')
        notary_block = generate_notary_block(state, county, name_of_individuals, type_of_authority, instrument_for, block_type='entity')
    # Replace placeholders in the document
    for paragraph in doc.paragraphs:
        if '[Signature Block]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Signature Block]', sig_block)
        if '[Signature Block With Notrary]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Signature Block With Notrary]', '')
        if '[Notary Block]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Notary Block]', notary_block)
    # Also replace in tables, headers, footers, and footnotes if needed
    def process_block(block):
        for paragraph in block.paragraphs:
            if '[Signature Block]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[Signature Block]', sig_block)
            if '[Signature Block With Notrary]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[Signature Block With Notrary]', '')
            if '[Notary Block]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[Notary Block]', notary_block)
        for table in getattr(block, 'tables', []):
            for row in table.rows:
                for cell in row.cells:
                    process_block(cell)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_block(cell)
    for section in doc.sections:
        process_block(section.header)
        process_block(section.footer)
    if hasattr(doc, 'part') and hasattr(doc.part, 'footnotes'):
        for footnote in doc.part.footnotes.part.footnotes:
            for paragraph in footnote.paragraphs:
                if '[Signature Block]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Signature Block]', sig_block)
                if '[Signature Block With Notrary]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Signature Block With Notrary]', '')
                if '[Notary Block]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Notary Block]', notary_block)
    return doc 

def build_exhibit_string(parcels):
    """
    Build the Exhibit A text string from parcel data (simplified version).
    
    Args:
        parcels: List of parcel objects with parcelNumber, isPortion, and templateType properties
    
    Returns:
        str: The complete Exhibit A text string
    """
    try:
        print(f"[DEBUG] Building exhibit string for {len(parcels)} parcels")
        
        # Validate parcels data
        if not isinstance(parcels, list) or len(parcels) == 0:
            raise ValueError("Parcels must be a non-empty list")
        
        # Start with header
        exhibit_parts = ["EXHIBIT A", "", "General Description of Property", ""]
        
        # Add image placeholder
        exhibit_parts.append("[Image]")
        exhibit_parts.append("")
        
        # Add parcel descriptions
        for i, parcel in enumerate(parcels, 1):
            if not isinstance(parcel, dict) or "parcelNumber" not in parcel:
                print(f"[WARNING] Invalid parcel data at index {i}: {parcel}")
                continue
            
            parcel_number = parcel.get("parcelNumber", i)
            is_portion = parcel.get("isPortion", False)
            
            # Simple template based on parcel type
            if is_portion:
                parcel_description = f"Portion {parcel_number}:\n\nThis portion of the property is described as follows: [Legal description for portion {parcel_number}]"
            else:
                parcel_description = f"Parcel {parcel_number}:\n\nA parcel of the property described as follows: [Legal description for parcel {parcel_number}]"
            
            print(f"[DEBUG] Parcel {parcel_number}: {{\"Portion\" if is_portion else \"Parcel\"}} (isPortion: {is_portion})")
            exhibit_parts.append(parcel_description)
            exhibit_parts.append("")  # Add spacing between parcels
        
        # Join all parts
        exhibit_string = "\n".join(exhibit_parts)
        
        print(f"[DEBUG] Generated exhibit string, length: {len(exhibit_string)}")
        return exhibit_string
        
    except Exception as e:
        print(f"[ERROR] Failed to build exhibit string: {str(e)}")
        import traceback
        traceback.print_exc()
        raise

# need to develop if condition to to stylstic 
def getNotaryBlock():
    """Get hardcoded notary block template"""
    return """STATE OF [State] SS: 

COUNTY OF [County] 

On _______________________, before me, __________________________________,  

Notary Public, personally appeared _______________________________________, who proved to me on the basis of satisfactory evidence to be the person(s) whose name(s) is/are subscribed to the within instrument and acknowledged to me that he/she/they executed the same in his/her/their authorized capacity(ies).] 

[STAMP]		________________________________ 

Title of Office: Notary Public 

Printed Name: ____________________ 

My Commission Expires: ___________"""

def getSigBlock(ownerType: str, numSignatures: int):
    # Store the values for future use
    owner_type = ownerType
    num_signatures = numSignatures
    filename1 = None
    filename2 = None
    filename1Content = None
    filename2Content = None
    # Map owner types to template files (using files that actually exist)
    if owner_type == 'his/her sole property' and num_signatures == 1:
        filename1 = 'SI1.txt'
    elif owner_type == 'a married couple' and num_signatures == 2:
        filename1 = 'I1.txt'
        filename2 = 'I1.txt'
    elif owner_type == 'Corporation':
        filename1 = 'E1.txt'
        if num_signatures == 2:
            filename2 = 'E1.txt'
    elif owner_type == 'LLC':
        filename1 = 'E1.txt'
        if num_signatures == 2:
            filename2 = 'E1.txt'
    elif owner_type == 'LP':    
        filename1 = 'E1.txt'
        if num_signatures == 2:
            filename2 = 'E1.txt'
    elif owner_type == 'Trust':
        filename1 = 'E1.txt'
        if num_signatures == 2:
            filename2 = 'E1.txt'
    elif owner_type == 'Sole Owner, married couple' and num_signatures == 2:
        filename1 = 'I1.txt'
        filename2 = 'SI1.txt'
    elif 'individual' in owner_type.lower():
        # Default individual case
        filename1 = 'I1.txt'
        if num_signatures == 2:
            filename2 = 'I1.txt'
    else:
        # Default entity case
        filename1 = 'E1.txt'
        if num_signatures == 2:
            filename2 = 'E1.txt'
    import os
    
    # Load content from filename1 if it exists
    if filename1:
        path1 = os.path.join('templates', 'sigBlocks', filename1)
        if os.path.exists(path1):
            with open(path1, 'r') as f:
                filename1Content = f.read().strip()
                print(f"[DEBUG] Loaded filename1 ({filename1}): {len(filename1Content)} characters")
        else:
            filename1Content = f"Template file '{filename1}' not found at {path1}"
            print(f"[ERROR] {filename1Content}")
    
    # Load content from filename2 if it exists  
    if filename2:
        path2 = os.path.join('templates', 'sigBlocks', filename2)
        if os.path.exists(path2):
            with open(path2, 'r') as f:
                filename2Content = f.read().strip()
                print(f"[DEBUG] Loaded filename2 ({filename2}): {len(filename2Content)} characters")
        else:
            filename2Content = f"Template file '{filename2}' not found at {path2}"
            print(f"[ERROR] {filename2Content}")
    
    print(f"[DEBUG] getSigBlock returning: filename1Content={filename1Content is not None}, filename2Content={filename2Content is not None}")
    
    # Return array with filename1 and filename2 content
    return [filename1Content, filename2Content]



# need to fix this function 
def notrary_generator():
    # Read notrary.txt file content
    import os
    notrary_file_path = os.path.join('templates', 'Notorary', 'notrary.txt')
    
    try:
        with open(notrary_file_path, 'r') as f:
            notrary_content = f.read()
        return notrary_content
    except FileNotFoundError:
        return f"Notary block template file 'notrary.txt' not found."
    except Exception as e:
        return f"Error reading notary block: {str(e)}"

# func desc
# sig_block is a txt file -> get sig block [file name]
# Notrary block is a txt file -> notrary generator [string]
# is_notray is boolean true or false [boolean]
# num_signatures is an integer [integer]


def generator(ownerType, is_notary, notary_block, num_signatures):
    owner_type = ownerType
    # Generate notary block with default parameters
    notary_content = notrary_generator()
    
    # Get signature block content
    filecontent = getSigBlock(owner_type, num_signatures)
    sigB1 = filecontent[0] if filecontent[0] is not None else ''
    sigB2 = filecontent[1] if filecontent[1] is not None else ''
    
    final_string = ""

    # Edge case where there are unique sig blocks 
    if owner_type == 'Sole owner, married couple' and is_notary:
        final_string += sigB1
        if is_notary and notary_content:
            final_string += "\n\n" + notary_content
        final_string += "\n\n" + sigB2
        if is_notary and notary_content:
            final_string += "\n\n" + notary_content

        return final_string
    if owner_type == 'Sole owner, married couple' and not is_notary:
        final_string += sigB1
        final_string += "\n\n" + sigB2
        return final_string
    
    # Generate additional signature blocks based on num_signatures
    for i in range(num_signatures):
        if is_notary and notary_content:
            final_string += "\n\n" + sigB1
            final_string += "\n\n" + notary_content
        else:
            final_string += "\n\n" + sigB1

    return final_string 



