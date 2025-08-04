"""
Enhanced image handling functionality for lease population
"""

import base64
import io
import os
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ImageEmbeddingHandler:
    """Enhanced image embedding with advanced features"""
    
    def __init__(self):
        self.supported_formats = {
            'PNG': b'\x89PNG\r\n\x1a\n',
            'JPEG': b'\xff\xd8\xff',
            'GIF': b'GIF87a',
            'GIF89a': b'GIF89a',
            'BMP': b'BM',
            'TIFF': b'II*\x00',  # Little-endian TIFF
            'TIFF_BE': b'MM\x00*',  # Big-endian TIFF
        }
        
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.max_width_inches = 6.0
        self.max_width_pixels = int(self.max_width_inches * 96)
        self.quality_settings = {
            'PNG': {'optimize': True, 'compress_level': 9},
            'JPEG': {'quality': 85, 'optimize': True},
            'WEBP': {'quality': 85, 'method': 6}
        }
    
    def validate_image_file(self, image_bytes: bytes) -> tuple[bool, str, str]:
        """
        Validate image file format and size
        
        Returns:
            tuple: (is_valid, format_name, error_message)
        """
        try:
            # Check file size
            if len(image_bytes) > self.max_file_size:
                return False, '', f"File too large (max {self.max_file_size // (1024*1024)}MB)"
            
            # Check minimum size
            if len(image_bytes) < 8:
                return False, '', "File too small to be a valid image"
            
            # Detect format
            for format_name, header in self.supported_formats.items():
                if image_bytes.startswith(header):
                    return True, format_name, ""
            
            return False, '', "Unsupported image format"
            
        except Exception as e:
            return False, '', f"Validation error: {str(e)}"
    
    def optimize_image(self, image: Image.Image, target_format: str = 'PNG') -> tuple[bytes, dict]:
        """
        Optimize image for document embedding
        
        Returns:
            tuple: (optimized_bytes, metadata)
        """
        try:
            metadata = {
                'original_size': image.size,
                'original_mode': image.mode,
                'format': target_format
            }
            
            # Convert to RGB if necessary
            if image.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                image = background
                metadata['converted_from'] = metadata['original_mode']
            elif image.mode != 'RGB':
                image = image.convert('RGB')
                metadata['converted_from'] = metadata['original_mode']
            
            # Resize if too large
            original_size = image.size
            if image.width > self.max_width_pixels:
                ratio = self.max_width_pixels / image.width
                new_width = self.max_width_pixels
                new_height = int(image.height * ratio)
                image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                metadata['resized_from'] = original_size
                metadata['resize_ratio'] = ratio
            
            # Save with optimization
            img_byte_arr = io.BytesIO()
            if target_format == 'PNG':
                image.save(img_byte_arr, format='PNG', **self.quality_settings['PNG'])
            elif target_format == 'JPEG':
                image.save(img_byte_arr, format='JPEG', **self.quality_settings['JPEG'])
            else:
                image.save(img_byte_arr, format='PNG', **self.quality_settings['PNG'])
            
            optimized_bytes = img_byte_arr.getvalue()
            metadata['final_size'] = image.size
            metadata['file_size'] = len(optimized_bytes)
            
            return optimized_bytes, metadata
            
        except Exception as e:
            logger.error(f"Image optimization failed: {str(e)}")
            raise
    
    def add_watermark(self, image: Image.Image, watermark_text: str = None) -> Image.Image:
        """
        Add optional watermark to image
        
        Args:
            image: PIL Image object
            watermark_text: Text to add as watermark
        
        Returns:
            Image with watermark
        """
        if not watermark_text:
            return image
        
        try:
            # Create a copy to avoid modifying original
            watermarked = image.copy()
            draw = ImageDraw.Draw(watermarked)
            
            # Try to load a font, fall back to default if not available
            try:
                font_size = max(12, min(image.width, image.height) // 20)
                font = ImageFont.truetype("arial.ttf", font_size)
            except:
                font = ImageFont.load_default()
            
            # Calculate text position (bottom right corner)
            bbox = draw.textbbox((0, 0), watermark_text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            
            x = image.width - text_width - 10
            y = image.height - text_height - 10
            
            # Add semi-transparent background
            bg_bbox = (x-5, y-5, x+text_width+5, y+text_height+5)
            draw.rectangle(bg_bbox, fill=(255, 255, 255, 128))
            
            # Add text
            draw.text((x, y), watermark_text, fill=(0, 0, 0, 128), font=font)
            
            return watermarked
            
        except Exception as e:
            logger.warning(f"Watermark addition failed: {str(e)}")
            return image
    
    def embed_image_enhanced(self, doc: Document, image_data: str, placeholder: str, 
                           watermark_text: str = None, target_format: str = 'PNG') -> dict:
        """
        Enhanced image embedding with comprehensive features
        
        Args:
            doc: DOCX document object
            image_data: Base64 encoded image string
            placeholder: Text placeholder to replace
            watermark_text: Optional watermark text
            target_format: Target format for optimization
        
        Returns:
            dict: Result with success status and metadata
        """
        result = {
            'success': False,
            'error': None,
            'metadata': {},
            'placeholders_found': 0
        }
        
        try:
            logger.info(f"Starting enhanced image embedding for placeholder: {placeholder}")
            
            # Validate input
            if not image_data or not isinstance(image_data, str):
                result['error'] = "Invalid image data: must be non-empty string"
                return result
            
            # Decode base64 image data
            try:
                image_bytes = base64.b64decode(image_data)
                logger.info(f"Decoded base64 image data, size: {len(image_bytes)} bytes")
            except Exception as e:
                result['error'] = f"Failed to decode base64 image data: {str(e)}"
                return result
            
            # Validate image format
            is_valid, format_name, error_msg = self.validate_image_file(image_bytes)
            if not is_valid:
                result['error'] = error_msg
                return result
            
            logger.info(f"Validated image format: {format_name}")
            
            # Process image with Pillow
            try:
                image = Image.open(io.BytesIO(image_bytes))
                logger.info(f"Opened image: format={image.format}, size={image.size}, mode={image.mode}")
                
                # Add watermark if specified
                if watermark_text:
                    image = self.add_watermark(image, watermark_text)
                    logger.info(f"Added watermark: {watermark_text}")
                
                # Optimize image
                optimized_bytes, metadata = self.optimize_image(image, target_format)
                result['metadata'] = metadata
                
                logger.info(f"Optimized image: {metadata}")
                
            except Exception as e:
                result['error'] = f"Failed to process image: {str(e)}"
                return result
            
            # Enhanced placeholder replacement
            found_placeholder = False
            placeholder_count = 0
            
            def process_paragraph(paragraph):
                nonlocal found_placeholder, placeholder_count
                if placeholder in paragraph.text:
                    placeholder_count += 1
                    logger.info(f"Found placeholder '{placeholder}' in paragraph #{placeholder_count}")
                    
                    # Clear paragraph and add centered image
                    paragraph.clear()
                    run = paragraph.add_run()
                    
                    # Calculate optimal image width
                    width_inches = min(metadata['final_size'][0] / 96, self.max_width_inches)
                    
                    # Add image with enhanced formatting
                    run.add_picture(io.BytesIO(optimized_bytes), width=Inches(width_inches))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add spacing after image
                    paragraph.space_after = Pt(12)
                    
                    found_placeholder = True
                    logger.info(f"Successfully embedded image in paragraph #{placeholder_count}")
            
            # Process all document sections
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
            
            # Process headers and footers
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    process_paragraph(paragraph)
                for paragraph in section.footer.paragraphs:
                    process_paragraph(paragraph)
            
            if not found_placeholder:
                result['error'] = f"Placeholder '{placeholder}' not found in document"
                return result
            
            result['success'] = True
            result['placeholders_found'] = placeholder_count
            logger.info(f"Enhanced image embedding completed successfully. Found {placeholder_count} placeholder(s)")
            
            return result
            
        except Exception as e:
            result['error'] = f"Critical error in enhanced image embedding: {str(e)}"
            logger.error(result['error'])
            import traceback
            logger.error(traceback.format_exc())
            return result
    
    def batch_process_images(self, doc: Document, image_mappings: list) -> dict:
        """
        Process multiple images in a single document
        
        Args:
            doc: DOCX document object
            image_mappings: List of dicts with 'placeholder', 'image_data', 'watermark' keys
        
        Returns:
            dict: Batch processing results
        """
        results = {
            'total_images': len(image_mappings),
            'successful': 0,
            'failed': 0,
            'results': []
        }
        
        for i, mapping in enumerate(image_mappings):
            try:
                result = self.embed_image_enhanced(
                    doc=doc,
                    image_data=mapping['image_data'],
                    placeholder=mapping['placeholder'],
                    watermark_text=mapping.get('watermark'),
                    target_format=mapping.get('format', 'PNG')
                )
                
                results['results'].append({
                    'index': i,
                    'placeholder': mapping['placeholder'],
                    'success': result['success'],
                    'error': result.get('error'),
                    'metadata': result.get('metadata', {})
                })
                
                if result['success']:
                    results['successful'] += 1
                else:
                    results['failed'] += 1
                    
            except Exception as e:
                results['results'].append({
                    'index': i,
                    'placeholder': mapping['placeholder'],
                    'success': False,
                    'error': str(e)
                })
                results['failed'] += 1
        
        return results 