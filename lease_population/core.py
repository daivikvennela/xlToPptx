"""
Core lease population processing functionality


Notes 
Having some issues with the image embedding. 
Need to break down the image embedding process into smaller steps. 
"""

import json
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from flask import jsonify, send_file
from .utils import normalize_placeholder_key, strip_brackets
from .image_handler import ImageEmbeddingHandler
from .block_replacer import embedImage, generate_signature_block, generate_notary_block, generator


class LeasePopulationProcessor:
    """
    Main processor for lease population functionality
    Handles document processing, placeholder replacement, and image embedding
    """
    
    def __init__(self):
        self.image_handler = ImageEmbeddingHandler()
    
    def _auto_generate_signature_blocks(self, mapping: dict) -> dict:
        """
        Auto-generate [Signature Block] and [Signature Block With Notrary] values
        based on detected party types and other relevant fields from the input mapping.
        """
        # Don't override if already provided
        if '[Signature Block]' in mapping and mapping['[Signature Block]'].strip():
            print("[INFO] [Signature Block] already provided, skipping auto-generation")
        if '[Signature Block With Notrary]' in mapping and mapping['[Signature Block With Notrary]'].strip():
            print("[INFO] [Signature Block With Notrary] already provided, skipping auto-generation")
            
        # If both are already provided, return as-is
        if ('[Signature Block]' in mapping and mapping['[Signature Block]'].strip() and
            '[Signature Block With Notrary]' in mapping and mapping['[Signature Block With Notrary]'].strip()):
            return mapping
        
        # Detect party type from mapping
        owner_type = self._detect_owner_type(mapping)
        print(f"[INFO] Detected owner type: {owner_type}")
        
        # Get number of signatures (default to 1)
        num_signatures = 1
        if '[Number of Grantor Signatures]' in mapping:
            try:
                num_signatures = int(mapping['[Number of Grantor Signatures]'])
            except (ValueError, TypeError):
                num_signatures = 1
        
        # Generate signature block without notary
        if '[Signature Block]' not in mapping or not mapping['[Signature Block]'].strip():
            try:
                signature_block = generator(owner_type, False, '', num_signatures)
                mapping['[Signature Block]'] = signature_block
                print(f"[INFO] Auto-generated [Signature Block] for {owner_type}")
            except Exception as e:
                print(f"[ERROR] Failed to generate signature block: {str(e)}")
                mapping['[Signature Block]'] = f"[Auto-generated signature block for {owner_type}]"
        
        # Generate signature block with notary
        if '[Signature Block With Notrary]' not in mapping or not mapping['[Signature Block With Notrary]'].strip():
            try:
                signature_with_notary = generator(owner_type, True, '', num_signatures)
                mapping['[Signature Block With Notrary]'] = signature_with_notary
                print(f"[INFO] Auto-generated [Signature Block With Notrary] for {owner_type}")
            except Exception as e:
                print(f"[ERROR] Failed to generate signature block with notary: {str(e)}")
                mapping['[Signature Block With Notrary]'] = f"[Auto-generated signature block with notary for {owner_type}]"
        
        return mapping
    
    def _detect_owner_type(self, mapping: dict) -> str:
        """
        Detect the owner type from the mapping based on available fields.
        Priority: [Owner Type] > [Grantor Type] > [Grantee Type] > default to 'individual'
        """
        # Check for explicit owner type
        for key in ['[Owner Type]', '[Grantor Type]', '[Grantee Type]']:
            if key in mapping and mapping[key].strip():
                owner_type = mapping[key].strip().lower()
                
                # Map variations to standard types
                if owner_type in ['individual', 'person']:
                    return 'individual'
                elif owner_type in ['entity', 'corporation', 'corp']:
                    return 'corporation'
                elif owner_type in ['llc', 'limited liability company']:
                    return 'llc'
                elif owner_type in ['lp', 'limited partnership']:
                    return 'lp'
                elif owner_type in ['married couple', 'couple']:
                    return 'married_couple'
                elif owner_type in ['sole owner married couple', 'sole owner, married couple']:
                    return 'Sole owner, married couple'
                else:
                    # Return the original value for any other type
                    return mapping[key].strip()
        
        # Default to individual if no type detected
        return 'individual'
    
    def process_lease_population(self, docx_file, mapping_json, track_changes=False, 
                               document_name='lease_population_filled', image_file=None):
        """
        Main processing function for lease population
        """
        try:
            # Parse mapping
            mapping = self._parse_mapping(mapping_json)
            
            # Auto-generate signature blocks if not provided
            mapping = self._auto_generate_signature_blocks(mapping)
            
            # Process image data
            mapping = self._process_image_data(mapping, image_file)
            
            # Load and process document
            doc = Document(docx_file)
            self._ensure_document_compatibility(doc)
            
            # Process image placeholders
            mapping = self._process_image_placeholders(doc, mapping)
            
            # Process text placeholders
            if track_changes:
                # Add "NEW:" prefix to all values for track changes
                mapping = {k: f"NEW:{v}" for k, v in mapping.items()}
                doc = self._replace_placeholders_with_track_changes(doc, mapping)
            else:
                doc = self._replace_placeholders_in_docx(doc, mapping)
            
            # Generate and return processed document
            return self._generate_final_document(doc, document_name)
            
        except Exception as e:
            import traceback
            error_traceback = traceback.format_exc()
            print(f"ERROR in lease_population_replace: {str(e)}")
            print(f"TRACEBACK: {error_traceback}")
            return jsonify({'error': f'Failed to process DOCX: {str(e)}', 'traceback': error_traceback}), 500
    
    def process_lease_population_enhanced(self, docx_file, mapping_json, track_changes=False, 
                                        document_name='lease_population_filled', image_files=None, 
                                        watermark_text=None, target_format='PNG'):
        """
        Enhanced processing function for lease population with advanced image support
        """
        try:
            # Parse mapping
            mapping = self._parse_mapping(mapping_json)
            
            # Auto-generate signature blocks if not provided
            mapping = self._auto_generate_signature_blocks(mapping)
            
            # Process multiple image files
            mapping = self._process_multiple_images(mapping, image_files, watermark_text, target_format)
            
            # Load and process document
            doc = Document(docx_file)
            self._ensure_document_compatibility(doc)
            
            # Process image placeholders with enhanced handler
            mapping = self._process_image_placeholders_enhanced(doc, mapping)
            
            # Process text placeholders
            if track_changes:
                # Add "NEW:" prefix to all values for track changes
                mapping = {k: f"NEW:{v}" for k, v in mapping.items()}
                doc = self._replace_placeholders_with_track_changes(doc, mapping)
            else:
                doc = self._replace_placeholders_in_docx(doc, mapping)
            
            # Generate and return processed document
            return self._generate_final_document(doc, document_name)
            
        except Exception as e:
            import traceback
            error_traceback = traceback.format_exc()
            print(f"ERROR in lease_population_replace_enhanced: {str(e)}")
            print(f"TRACEBACK: {error_traceback}")
            return jsonify({'error': f'Failed to process DOCX: {str(e)}', 'traceback': error_traceback}), 500
    
    def _parse_mapping(self, mapping_json):
        """Parse and validate mapping JSON"""
        if not mapping_json or mapping_json.strip() in ['undefined', 'null']:
            raise ValueError('No key-value mapping provided')
        
        mapping_raw = json.loads(mapping_json)
        if not mapping_raw or not isinstance(mapping_raw, list):
            raise ValueError('Invalid mapping format')
        
        return {item['key']: item['value'] for item in mapping_raw if item['value'].strip()}
    
    def _process_image_data(self, mapping, image_file):
        """Process image file data and add to mapping"""
        exhibit_a_image_1 = None
        
        if image_file:
            try:
                image_data = image_file.read()
                exhibit_a_image_1 = base64.b64encode(image_data).decode('utf-8')
                print(f"[DEBUG] Processed image file, size: {len(exhibit_a_image_1)} characters")
            except Exception as e:
                print(f"[ERROR] Failed to process image file: {str(e)}")
        
        if exhibit_a_image_1:
            mapping['[EXHIBIT_A_IMAGE_1]'] = exhibit_a_image_1
        
        return mapping
    
    def _process_multiple_images(self, mapping, image_files, watermark_text=None, target_format='PNG'):
        """Process multiple image files with enhanced features"""
        if not image_files:
            return mapping
        
        try:
            for key, image_file in image_files.items():
                try:
                    image_data = image_file.read()
                    image_b64 = base64.b64encode(image_data).decode('utf-8')
                    
                    # Determine placeholder based on file key
                    if key == 'exhibit_image':
                        placeholder = '[EXHIBIT_A_IMAGE_1]'
                    elif key.startswith('image_'):
                        placeholder = key.replace('image_', '[').upper() + ']'
                    else:
                        placeholder = f'[{key.upper()}]'
                    
                    # Store image data with metadata
                    mapping[placeholder] = {
                        'image_data': image_b64,
                        'watermark': watermark_text,
                        'format': target_format,
                        'original_filename': image_file.filename,
                        'size': len(image_data)
                    }
                    
                    print(f"Processed image {key}: {image_file.filename} -> {placeholder}")
                    
                except Exception as e:
                    print(f"Failed to process image {key}: {str(e)}")
                    continue
            
            return mapping
            
        except Exception as e:
            print(f"Error processing multiple images: {str(e)}")
            return mapping
    
    def _process_image_placeholders_enhanced(self, doc, mapping):
        """Process image placeholders with enhanced image handler"""
        image_placeholders = []
        
        for key, value in mapping.items():
            # Check if this is an image placeholder with enhanced data
            if isinstance(value, dict) and 'image_data' in value:
                image_placeholders.append((key, value))
                print(f"Found enhanced image placeholder: {key}")
            elif isinstance(value, str) and value.strip() and key.strip().lower() == '[image]':
                # Legacy image handling
                image_placeholders.append((key, {'image_data': value}))
                print(f"Found legacy image placeholder: {key}")
        
        # Handle image embedding for each image placeholder
        for placeholder_key, image_config in image_placeholders:
            try:
                print(f"Attempting to embed image for placeholder: {placeholder_key}")
                
                # Use enhanced image handler
                result = self.image_handler.embed_image_enhanced(
                    doc=doc,
                    image_data=image_config['image_data'],
                    placeholder=placeholder_key,
                    watermark_text=image_config.get('watermark'),
                    target_format=image_config.get('format', 'PNG')
                )
                
                if result['success']:
                    print(f"Image embedding successful for {placeholder_key}")
                    print(f"Metadata: {result.get('metadata', {})}")
                    # Clear value to prevent text replacement
                    mapping[placeholder_key] = ''
                else:
                    print(f"Image embedding failed for {placeholder_key}: {result.get('error')}")
                    
            except Exception as e:
                print(f"Image embedding error for {placeholder_key}: {str(e)}")
                import traceback
                traceback.print_exc()
        
        return mapping
    
    def _ensure_document_compatibility(self, doc):
        """Ensure document compatibility with modern Word"""
        if hasattr(doc.core_properties, 'version'):
            doc.core_properties.version = '16.0'
        if hasattr(doc.core_properties, 'last_modified_by'):
            doc.core_properties.last_modified_by = 'Document Processor'
    
    def _replace_placeholders_in_docx(self, doc, mapping):
        """Replace placeholders in DOCX document"""
        def replace_in_runs(runs, mapping):
            full_text = ''.join(run.text for run in runs)
            for key, value in mapping.items():
                if not value.strip():
                    continue
                for variant in normalize_placeholder_key(key):
                    full_text = full_text.replace(variant, value)
                    bracketless = strip_brackets(variant)
                    if bracketless != variant:
                        full_text = full_text.replace(bracketless, value)
            if runs:
                runs[0].text = full_text
                for run in runs[1:]:
                    run.text = ''
        
        def process_paragraph(paragraph, mapping):
            if not paragraph.runs:
                return
            joined = ''.join(run.text for run in paragraph.runs)
            if any(variant in joined for key, value in mapping.items() if value.strip() 
                   for variant in normalize_placeholder_key(key)):
                replace_in_runs(paragraph.runs, mapping)
        
        def process_table(table, mapping):
            for row in table.rows:
                for cell in row.cells:
                    process_block(cell, mapping)
        
        def process_block(block, mapping):
            for paragraph in block.paragraphs:
                process_paragraph(paragraph, mapping)
            for table in getattr(block, 'tables', []):
                process_table(table, mapping)
        
        # Process all document sections
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph, mapping)
        for table in doc.tables:
            process_table(table, mapping)
        for section in doc.sections:
            process_block(section.header, mapping)
            process_block(section.footer, mapping)
        if hasattr(doc, 'part') and hasattr(doc.part, 'footnotes'):
            for footnote in doc.part.footnotes.part.footnotes:
                for paragraph in footnote.paragraphs:
                    process_paragraph(paragraph, mapping)
        
        return doc
    
    def _replace_placeholders_with_track_changes(self, doc, mapping):
        """Replace placeholders with track changes highlighting"""
        def process_paragraph(paragraph, mapping):
            for run in paragraph.runs:
                for key, value in mapping.items():
                    if not value.strip():
                        continue
                    replaced = False
                    for variant in normalize_placeholder_key(key):
                        if variant in run.text:
                            run.text = run.text.replace(variant, value)
                            run.font.highlight_color = 7  # yellow
                            replaced = True
                            break
                    if replaced:
                        break
        
        def process_table(table, mapping):
            for row in table.rows:
                for cell in row.cells:
                    process_block(cell, mapping)
        
        def process_block(block, mapping):
            for paragraph in block.paragraphs:
                process_paragraph(paragraph, mapping)
            for table in getattr(block, 'tables', []):
                process_table(table, mapping)
        
        # Process all document sections
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph, mapping)
        for table in doc.tables:
            process_table(table, mapping)
        for section in doc.sections:
            process_block(section.header, mapping)
            process_block(section.footer, mapping)
        if hasattr(doc, 'part') and hasattr(doc.part, 'footnotes'):
            for footnote in doc.part.footnotes.part.footnotes:
                for paragraph in footnote.paragraphs:
                    process_paragraph(paragraph, mapping)
        
        return doc
    
    def _generate_final_document(self, doc, document_name):
        """Generate final DOCX document for download"""
        out_stream = BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)
        safe_name = document_name.replace(' ', '_').replace('/', '_')
        print(f"[DEBUG] lease_population_replace: Download filename will be: {safe_name}.docx")
        return send_file(out_stream, as_attachment=True, download_name=f'{safe_name}.docx', 
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    def test_party_type(self, docx_file, mapping_json, party_type, document_name='party_type_test'):
        """Test party type functionality"""
        try:
            mapping = self._parse_mapping(mapping_json)
            mapping['[Grantor Type]'] = party_type
            
            doc = Document(docx_file)
            self._ensure_document_compatibility(doc)
            
            # Prepare values for template filling
            grantor_name = mapping.get('[Grantor Name]', '')
            trust_entity_name = mapping.get('[Trust/Entity Name]', '')
            name = mapping.get('[Name]', '')
            title = mapping.get('[Title]', '')
            state = mapping.get('[State]', '')
            county = mapping.get('[County]', '')
            name_of_individuals = mapping.get('[NAME(S) OF INDIVIDUAL(S)]', '')
            type_of_authority = mapping.get('[TYPE OF AUTHORITY]', '')
            instrument_for = mapping.get('[NAME OF ENTITY OR TRUST WHOM INSTRUMENT WAS EXECUTED FOR]', '')
            
            is_individual = party_type.lower() == 'individual'
            if is_individual:
                sig_block = generate_signature_block(grantor_name, block_type='individual')
                notary_block = generate_notary_block(state, county, name_of_individuals, block_type='individual')
            else:
                sig_block = generate_signature_block(grantor_name, trust_entity_name, name, title, block_type='entity')
                notary_block = generate_notary_block(state, county, name_of_individuals, type_of_authority, instrument_for, block_type='entity')
            
            # Replacement logic
            def replace_blocks_in_runs(runs):
                for run in runs:
                    if '[Signature Block]' in run.text:
                        run.text = run.text.replace('[Signature Block]', sig_block)
                    if '[Notary Block]' in run.text:
                        run.text = run.text.replace('[Notary Block]', notary_block)
            
            def process_paragraph(paragraph):
                replace_blocks_in_runs(paragraph.runs)
            
            def process_table(table):
                for row in table.rows:
                    for cell in row.cells:
                        process_block(cell)
            
            def process_block(block):
                for paragraph in block.paragraphs:
                    process_paragraph(paragraph)
                for table in getattr(block, 'tables', []):
                    process_table(table)
            
            # Process all document sections
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)
            for table in doc.tables:
                process_table(table)
            for section in doc.sections:
                process_block(section.header)
                process_block(section.footer)
            if hasattr(doc, 'part') and hasattr(doc.part, 'footnotes'):
                for footnote in doc.part.footnotes.part.footnotes:
                    for paragraph in footnote.paragraphs:
                        process_paragraph(paragraph)
            
            return self._generate_final_document(doc, document_name)
            
        except Exception as e:
            import traceback
            error_traceback = traceback.format_exc()
            print(f"ERROR in test_party_type: {str(e)}")
            print(f"TRACEBACK: {error_traceback}")
            return jsonify({'error': f'Failed to process DOCX: {str(e)}', 'traceback': error_traceback}), 500 