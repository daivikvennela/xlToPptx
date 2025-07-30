import os
import json
import logging
from typing import Dict, List, Optional, Tuple
from io import BytesIO
import tempfile
import shutil

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement
    from docx.shared import Inches
    import xml.etree.ElementTree as ET
    from lxml import etree
except ImportError as e:
    print(f"Warning: Required libraries not available: {e}")
    Document = None

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FootnoteEditor:
    """
    A comprehensive footnote extraction and editing tool for DOCX files.
    Uses python-docx for reliable footnote access with XML fallback.
    """
    
    def __init__(self):
        self.supported_extensions = ['.docx']
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
        
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """
        Validate uploaded file for footnote processing.
        
        Args:
            file_path: Path to the uploaded file
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            # Check file exists
            if not os.path.exists(file_path):
                return False, "File not found"
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                return False, f"File too large ({file_size / 1024 / 1024:.1f}MB). Max size: {self.max_file_size / 1024 / 1024}MB"
            
            # Check file extension
            _, ext = os.path.splitext(file_path)
            if ext.lower() not in self.supported_extensions:
                return False, f"Unsupported file type: {ext}. Supported: {', '.join(self.supported_extensions)}"
            
            return True, ""
            
        except Exception as e:
            logger.error(f"File validation error: {str(e)}")
            return False, f"File validation failed: {str(e)}"
    
    def extract_footnotes(self, docx_path: str) -> Dict:
        """
        Extract all footnotes from a DOCX file.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dictionary containing footnote data and metadata
        """
        try:
            if Document is None:
                return {"error": "python-docx library not available"}
            
            doc = Document(docx_path)
            
            # Initialize result structure
            result = {
                "success": True,
                "footnotes": [],
                "metadata": {
                    "total_footnotes": 0,
                    "has_footnotes": False,
                    "file_name": os.path.basename(docx_path)
                }
            }
            
            # Check if document has footnotes
            if not hasattr(doc, 'footnotes') or doc.footnotes is None:
                result["metadata"]["has_footnotes"] = False
                return result
            
            # Extract footnotes using python-docx API
            footnotes = doc.footnotes._footnotes if hasattr(doc.footnotes, '_footnotes') else []
            
            if not footnotes:
                # Try alternative access method
                try:
                    footnotes = list(doc.footnotes)
                except Exception as e:
                    logger.warning(f"Could not access footnotes via API: {e}")
                    footnotes = []
            
            result["metadata"]["total_footnotes"] = len(footnotes)
            result["metadata"]["has_footnotes"] = len(footnotes) > 0
            
            # Extract footnote content
            for i, footnote in enumerate(footnotes):
                try:
                    footnote_data = {
                        "id": i + 1,
                        "original_text": "",
                        "edited_text": "",
                        "has_changes": False,
                        "paragraphs": []
                    }
                    
                    # Extract text from all paragraphs in the footnote
                    if hasattr(footnote, 'paragraphs'):
                        for para in footnote.paragraphs:
                            para_text = para.text.strip()
                            if para_text:
                                footnote_data["paragraphs"].append(para_text)
                    
                    # Combine all paragraph text
                    footnote_data["original_text"] = "\n".join(footnote_data["paragraphs"])
                    footnote_data["edited_text"] = footnote_data["original_text"]
                    
                    result["footnotes"].append(footnote_data)
                    
                except Exception as e:
                    logger.error(f"Error extracting footnote {i+1}: {e}")
                    # Add error footnote entry
                    result["footnotes"].append({
                        "id": i + 1,
                        "original_text": f"[ERROR: Could not extract footnote {i+1}]",
                        "edited_text": f"[ERROR: Could not extract footnote {i+1}]",
                        "has_changes": False,
                        "paragraphs": [],
                        "error": str(e)
                    })
            
            logger.info(f"Successfully extracted {len(result['footnotes'])} footnotes from {docx_path}")
            return result
            
        except Exception as e:
            logger.error(f"Error extracting footnotes: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to extract footnotes: {str(e)}",
                "footnotes": [],
                "metadata": {
                    "total_footnotes": 0,
                    "has_footnotes": False,
                    "file_name": os.path.basename(docx_path)
                }
            }
    
    def update_footnotes(self, docx_path: str, footnotes_data: List[Dict]) -> Dict:
        """
        Update footnotes in a DOCX file with edited content.
        
        Args:
            docx_path: Path to the original DOCX file
            footnotes_data: List of footnote dictionaries with edited text
            
        Returns:
            Dictionary with success status and output file path
        """
        try:
            if Document is None:
                return {"error": "python-docx library not available"}
            
            # Load the document
            doc = Document(docx_path)
            
            # Check if document has footnotes
            if not hasattr(doc, 'footnotes') or doc.footnotes is None:
                return {
                    "success": False,
                    "error": "Document has no footnotes to update"
                }
            
            # Get footnotes
            footnotes = doc.footnotes._footnotes if hasattr(doc.footnotes, '_footnotes') else []
            if not footnotes:
                try:
                    footnotes = list(doc.footnotes)
                except Exception as e:
                    logger.warning(f"Could not access footnotes via API: {e}")
                    return {
                        "success": False,
                        "error": "Could not access footnotes in document"
                    }
            
            # Update footnotes with edited content
            updated_count = 0
            errors = []
            
            for i, footnote in enumerate(footnotes):
                if i >= len(footnotes_data):
                    break
                
                try:
                    footnote_info = footnotes_data[i]
                    
                    # Check if this footnote has changes
                    if footnote_info.get("has_changes", False):
                        edited_text = footnote_info.get("edited_text", "")
                        
                        # Clear existing content
                        if hasattr(footnote, 'paragraphs'):
                            # Remove all paragraphs except the first one
                            while len(footnote.paragraphs) > 1:
                                footnote.paragraphs[-1]._element.getparent().remove(footnote.paragraphs[-1]._element)
                            
                            # Update the first paragraph
                            if footnote.paragraphs:
                                footnote.paragraphs[0].text = edited_text
                            else:
                                # Add a new paragraph if none exists
                                footnote.add_paragraph(edited_text)
                        
                        updated_count += 1
                        logger.info(f"Updated footnote {i+1}")
                    
                except Exception as e:
                    error_msg = f"Error updating footnote {i+1}: {str(e)}"
                    logger.error(error_msg)
                    errors.append(error_msg)
            
            # Create output file
            output_path = self._create_output_path(docx_path)
            
            # Save the modified document
            doc.save(output_path)
            
            result = {
                "success": True,
                "output_path": output_path,
                "updated_count": updated_count,
                "total_footnotes": len(footnotes),
                "errors": errors
            }
            
            logger.info(f"Successfully updated {updated_count} footnotes in {output_path}")
            return result
            
        except Exception as e:
            logger.error(f"Error updating footnotes: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to update footnotes: {str(e)}"
            }
    
    def batch_replace_footnotes(self, docx_path: str, search_text: str, replace_text: str) -> Dict:
        """
        Perform batch search and replace across all footnotes.
        
        Args:
            docx_path: Path to the DOCX file
            search_text: Text to search for
            replace_text: Text to replace with
            
        Returns:
            Dictionary with replacement results
        """
        try:
            # First extract footnotes
            extraction_result = self.extract_footnotes(docx_path)
            
            if not extraction_result.get("success", False):
                return extraction_result
            
            footnotes = extraction_result["footnotes"]
            replacement_count = 0
            
            # Perform batch replacement
            for footnote in footnotes:
                original_text = footnote.get("original_text", "")
                if search_text in original_text:
                    new_text = original_text.replace(search_text, replace_text)
                    footnote["edited_text"] = new_text
                    footnote["has_changes"] = True
                    replacement_count += 1
            
            # Update the document with changes
            if replacement_count > 0:
                update_result = self.update_footnotes(docx_path, footnotes)
                update_result["replacement_count"] = replacement_count
                update_result["search_text"] = search_text
                update_result["replace_text"] = replace_text
                return update_result
            else:
                return {
                    "success": True,
                    "message": f"No occurrences of '{search_text}' found in footnotes",
                    "replacement_count": 0
                }
                
        except Exception as e:
            logger.error(f"Error in batch replace: {str(e)}")
            return {
                "success": False,
                "error": f"Batch replace failed: {str(e)}"
            }
    
    def _create_output_path(self, original_path: str) -> str:
        """
        Create a unique output path for the modified document.
        
        Args:
            original_path: Path to the original document
            
        Returns:
            Path for the output document
        """
        import tempfile
        import datetime
        
        # Create uploads directory if it doesn't exist
        uploads_dir = "uploads"
        os.makedirs(uploads_dir, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = os.path.splitext(os.path.basename(original_path))[0]
        output_filename = f"{base_name}_footnotes_edited_{timestamp}.docx"
        
        return os.path.join(uploads_dir, output_filename)
    
    def get_footnote_statistics(self, docx_path: str) -> Dict:
        """
        Get statistics about footnotes in a document.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dictionary with footnote statistics
        """
        try:
            extraction_result = self.extract_footnotes(docx_path)
            
            if not extraction_result.get("success", False):
                return extraction_result
            
            footnotes = extraction_result["footnotes"]
            
            # Calculate statistics
            total_chars = sum(len(f.get("original_text", "")) for f in footnotes)
            total_words = sum(len(f.get("original_text", "").split()) for f in footnotes)
            empty_footnotes = sum(1 for f in footnotes if not f.get("original_text", "").strip())
            
            stats = {
                "success": True,
                "total_footnotes": len(footnotes),
                "total_characters": total_chars,
                "total_words": total_words,
                "empty_footnotes": empty_footnotes,
                "average_chars_per_footnote": total_chars / len(footnotes) if footnotes else 0,
                "average_words_per_footnote": total_words / len(footnotes) if footnotes else 0
            }
            
            return stats
            
        except Exception as e:
            logger.error(f"Error getting footnote statistics: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to get statistics: {str(e)}"
            }

# Global instance for use in Flask routes
footnote_editor = FootnoteEditor() 